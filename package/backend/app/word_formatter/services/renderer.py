"""
确定性渲染器：AST + reference.docx → output.docx
"""
from __future__ import annotations

import io
import os
import re
from dataclasses import dataclass
from typing import Optional, Set

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm

from ..models.ast import (
    BibliographyBlock,
    CodeBlock,
    DocumentAST,
    FigureBlock,
    HeadingBlock,
    Inline,
    ListBlock,
    ParagraphBlock,
    PageBreakBlock,
    SectionBreakBlock,
    TableBlock,
)
from ..models.stylespec import StyleSpec


def _align_to_docx(align: str):
    return {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }[align]


def _clear_paragraph_runs(p) -> None:
    """清除段落中的所有 runs。"""
    for r in list(p.runs):
        try:
            p._p.remove(r._r)
        except Exception:
            pass


def _clear_paragraph_numbering(p) -> None:
    """清除段落的自动编号格式。

    Word 的标题样式可能关联了 Outline Numbering，
    这会导致标题前出现额外的 1, 2, 3 等序号。
    通过移除 numPr 元素可以清除这些自动编号。
    """
    pPr = p._p.find(qn("w:pPr"))
    if pPr is not None:
        numPr = pPr.find(qn("w:numPr"))
        if numPr is not None:
            pPr.remove(numPr)


def _apply_inline_style(run, inline_type: str) -> None:
    """根据 Inline 类型应用对应的样式到 run。"""
    if inline_type == "bold":
        run.bold = True
    elif inline_type == "italic":
        run.italic = True
    elif inline_type == "underline":
        run.underline = True
    elif inline_type == "superscript":
        run.font.superscript = True
    elif inline_type == "subscript":
        run.font.subscript = True
    elif inline_type == "code":
        run.font.name = "Consolas"


def _apply_inlines(p, inlines: list) -> None:
    """
    将 Inline 列表渲染到段落中，保留富文本格式。

    Args:
        p: python-docx 段落对象
        inlines: Inline 对象列表
    """
    for inline in inlines:
        text = inline.text or ""
        # 处理换行符
        parts = text.split("\n")
        for i, part in enumerate(parts):
            if i > 0:
                p.add_run().add_break()
            if not part:
                continue
            run = p.add_run(part)
            _apply_inline_style(run, inline.type)


@dataclass
class RenderOptions:
    include_cover: bool = True
    include_toc: bool = True
    toc_title: str = "目 录"
    toc_levels: int = 3


_FRONT_HEADINGS: Set[str] = {
    "摘要", "关键词", "关键字", "abstract", "key words", "keywords",
    "致谢", "谢辞", "参考文献", "references", "目录", "目 录",
}

_FRONT_ONLY_HEADINGS: Set[str] = {
    "摘要", "关键词", "关键字", "abstract", "key words", "keywords",
}


def _is_front_heading(text: str) -> bool:
    """检查是否为前置标题（不区分大小写）。"""
    return text.lower() in _FRONT_HEADINGS or text in _FRONT_HEADINGS


def _is_front_only_heading(text: str) -> bool:
    """检查是否为仅前置标题（不区分大小写）。"""
    return text.lower() in _FRONT_ONLY_HEADINGS or text in _FRONT_ONLY_HEADINGS


def _apply_page_numbering_ooxml(docx_bytes: bytes, spec: StyleSpec) -> bytes:
    """Set section-based page numbering format/start using OOXML."""
    pn = spec.page_numbering
    if not pn or not pn.enabled:
        return docx_bytes
    from ..utils.ooxml import DocxPackage
    from lxml import etree

    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    NSMAP = {"w": W_NS}

    def _qn(tag: str) -> str:
        pref, local = tag.split(":")
        return f"{{{NSMAP[pref]}}}{local}"

    fmt_map = {
        "decimal": "decimal",
        "romanUpper": "upperRoman",
        "romanLower": "lowerRoman",
    }

    pkg = DocxPackage.from_bytes(docx_bytes)
    root = pkg.read_xml("word/document.xml")
    body = root.find("w:body", namespaces=NSMAP)
    if body is None:
        return docx_bytes

    sect_prs = body.findall(".//w:sectPr", namespaces=NSMAP)
    seen = set()
    ordered = []
    for s in sect_prs:
        sid = id(s)
        if sid in seen:
            continue
        seen.add(sid)
        ordered.append(s)
    if not ordered:
        return docx_bytes

    def _set_pgnum(sectPr, fmt: str, start: int):
        pg = sectPr.find("w:pgNumType", namespaces=NSMAP)
        if pg is None:
            pg = etree.SubElement(sectPr, _qn("w:pgNumType"))
        pg.set(_qn("w:fmt"), fmt_map.get(fmt, "decimal"))
        pg.set(_qn("w:start"), str(int(start)))

    if len(ordered) == 1:
        _set_pgnum(ordered[0], pn.main_format, pn.main_start)
    else:
        _set_pgnum(ordered[0], pn.front_format, pn.front_start)
        _set_pgnum(ordered[-1], pn.main_format, pn.main_start)

    pkg.write_xml("word/document.xml", root)
    return pkg.to_bytes()


def _insert_toc_paragraph(doc: Document, title: str, front_style: str, max_level: int):
    doc.add_paragraph("")
    doc.add_paragraph("")
    p = doc.add_paragraph(title)
    if front_style in doc.styles:
        p.style = doc.styles[front_style]
    p2 = doc.add_paragraph()
    run = p2.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), f'TOC \\o "1-{max_level}" \\h \\z \\u')
    run._r.append(fld)
    run2 = p2.add_run("（在 Word 中右键目录 → 更新域）")


def _detect_heading_level_offset(ast: DocumentAST) -> int:
    """检测标题级别偏移量。

    如果文档中最小的标题级别不是 1（例如用户用 ## 作为一级标题），
    计算需要减去的偏移量，使其归一化为从 1 开始。

    Args:
        ast: 文档 AST

    Returns:
        偏移量（0 表示无需调整，1 表示所有级别减 1，以此类推）
    """
    min_level = None
    for block in ast.blocks:
        if isinstance(block, HeadingBlock):
            # 跳过前置标题（摘要、Abstract 等通常是单独设置的）
            heading_text = block.text.strip().lower()
            if heading_text in _FRONT_HEADINGS:
                continue
            if min_level is None or block.level < min_level:
                min_level = block.level

    # 如果没有正文标题或最小级别已是 1，无需调整
    if min_level is None or min_level <= 1:
        return 0

    return min_level - 1


def render_docx(
    ast: DocumentAST,
    spec: StyleSpec,
    reference_docx_bytes: bytes,
    options: Optional[RenderOptions] = None,
) -> bytes:
    options = options or RenderOptions()
    doc = Document(io.BytesIO(reference_docx_bytes))

    section = doc.sections[0]
    section.top_margin = Mm(spec.page.margins_mm.top)
    section.bottom_margin = Mm(spec.page.margins_mm.bottom)
    section.left_margin = Mm(spec.page.margins_mm.left)
    section.right_margin = Mm(spec.page.margins_mm.right)
    section.gutter = Mm(spec.page.margins_mm.binding)
    section.header_distance = Mm(spec.page.header_mm)
    section.footer_distance = Mm(spec.page.footer_mm)

    # clear initial dummy paragraph if empty
    if len(doc.paragraphs) == 1 and not doc.paragraphs[0].text.strip():
        p = doc.paragraphs[0]._element
        p.getparent().remove(p)

    # Cover
    if options.include_cover:
        _render_cover(doc, ast)
        doc.add_page_break()

    # TOC
    if options.include_toc:
        _insert_toc_paragraph(doc, options.toc_title, "FrontHeading", spec.structure.toc_max_level)
        doc.add_page_break()

    need_page_numbering = bool(spec.page_numbering and spec.page_numbering.enabled)
    main_section_inserted = False

    current_section = None
    abstract_prefixed = False
    keywords_prefixed = False
    fig_counter = 0
    table_counter = 0

    # 检测标题级别偏移（支持 ## 作为一级标题等情况）
    heading_level_offset = _detect_heading_level_offset(ast)

    for block in ast.blocks:
        if isinstance(block, HeadingBlock):
            heading_text = block.text.strip()

            if heading_text in {"摘要"}:
                current_section = "cn_abstract"
            elif heading_text in {"关键词", "关键字"}:
                current_section = "cn_keywords"
            elif heading_text.lower() == "abstract":
                current_section = "en_abstract"
            elif heading_text.lower() in {"key words", "keywords"}:
                current_section = "en_keywords"
            elif heading_text.lower() in {"参考文献", "references"}:
                current_section = "references"
            else:
                current_section = "body"

            if (
                need_page_numbering
                and not main_section_inserted
                and not _is_front_only_heading(heading_text)
                and len(doc.paragraphs) > 0
            ):
                doc.add_section(WD_SECTION.NEW_PAGE)
                main_section_inserted = True

            if _is_front_heading(heading_text):
                style_id = "FrontHeading"
                display_text = heading_text
            else:
                # 应用级别偏移（支持 ## 作为一级标题等情况）
                effective_level = max(1, block.level - heading_level_offset)

                # 限制最大层级为 3（模板通常只支持 H1-H3）
                effective_level = min(effective_level, 3)

                if effective_level == 1:
                    style_id = "H1"
                elif effective_level == 2:
                    style_id = "H2"
                elif effective_level == 3:
                    style_id = "H3"
                else:
                    style_id = "H3"

                # 直接使用原始标题文本，不添加编号
                display_text = heading_text

            p = doc.add_paragraph(display_text)
            if style_id in doc.styles:
                p.style = doc.styles[style_id]
            elif "Body" in doc.styles:
                p.style = doc.styles["Body"]
            # 清除 Word 样式可能关联的自动编号
            _clear_paragraph_numbering(p)
            continue

        if isinstance(block, ParagraphBlock):
            inlines = block.inlines
            raw_text = block.text
            if raw_text is None and inlines:
                raw_text = "".join(i.text for i in inlines)
            if not (raw_text or "").strip():
                continue

            # 处理摘要/关键词前缀
            if spec.auto_prefix_abstract_keywords:
                if current_section == "cn_abstract" and not abstract_prefixed:
                    if not (raw_text or "").startswith("摘要："):
                        if inlines:
                            inlines = [Inline(type="text", text="摘要：")] + list(inlines)
                        else:
                            raw_text = "摘要：" + (raw_text or "")
                    abstract_prefixed = True
                elif current_section == "en_abstract" and not abstract_prefixed:
                    if not (raw_text or "").lower().startswith("abstract:"):
                        if inlines:
                            inlines = [Inline(type="text", text="Abstract: ")] + list(inlines)
                        else:
                            raw_text = "Abstract: " + (raw_text or "")
                    abstract_prefixed = True
                elif current_section in {"cn_keywords", "en_keywords"} and not keywords_prefixed:
                    if current_section == "cn_keywords" and not (raw_text or "").startswith(("关键词：", "关键字：")):
                        if inlines:
                            inlines = [Inline(type="text", text="关键词：")] + list(inlines)
                        else:
                            raw_text = "关键词：" + _normalize_cn_keywords(raw_text or "")
                    elif current_section == "en_keywords" and not (raw_text or "").lower().startswith(("key words:", "keywords:")):
                        if inlines:
                            inlines = [Inline(type="text", text="Key words: ")] + list(inlines)
                        else:
                            raw_text = "Key words: " + _normalize_en_keywords(raw_text or "")
                    keywords_prefixed = True

            style_id = "Body"
            if current_section in {"cn_abstract", "en_abstract"}:
                style_id = "AbstractBody"
            elif current_section in {"cn_keywords", "en_keywords"}:
                style_id = "KeywordsBody"
            elif current_section == "references":
                style_id = "Reference"

            p = doc.add_paragraph()
            if style_id in doc.styles:
                p.style = doc.styles[style_id]
            # 优先使用富文本渲染
            if inlines:
                _apply_inlines(p, inlines)
            else:
                p.add_run(raw_text or "")
            continue

        if isinstance(block, ListBlock):
            style_name = "ListNumber" if block.ordered else "ListBullet"
            use_style = style_name in doc.styles
            for idx, item in enumerate(block.items, start=1):
                raw_text = "".join(i.text for i in item.inlines)
                if not raw_text.strip():
                    continue
                if use_style:
                    p = doc.add_paragraph()
                    p.style = doc.styles[style_name]
                    _apply_inlines(p, item.inlines)
                else:
                    prefix = f"{idx}. " if block.ordered else "• "
                    p = doc.add_paragraph()
                    if "Body" in doc.styles:
                        p.style = doc.styles["Body"]
                    p.add_run(prefix)
                    _apply_inlines(p, item.inlines)
            continue

        if isinstance(block, TableBlock):
            if block.caption:
                caption = block.caption.strip()
                if spec.auto_number_figures_tables and not re.match(r"^表\d+", caption):
                    table_counter += 1
                    caption = f"表{table_counter} {caption}"
                pcap = doc.add_paragraph(caption)
                if "TableTitle" in doc.styles:
                    pcap.style = doc.styles["TableTitle"]
            elif spec.auto_number_figures_tables:
                table_counter += 1
                pcap = doc.add_paragraph(f"表{table_counter}")
                if "TableTitle" in doc.styles:
                    pcap.style = doc.styles["TableTitle"]
            if not block.rows:
                continue
            # 使用富文本列或纯文本列来计算列数
            rows_for_cols = block.rows_inlines if block.rows_inlines else block.rows
            cols = max(len(r) for r in rows_for_cols)
            table = doc.add_table(rows=len(block.rows), cols=cols)
            for r_i, row in enumerate(block.rows):
                for c_i in range(cols):
                    cell = table.cell(r_i, c_i)
                    # 优先使用富文本渲染
                    cell_inlines = None
                    if block.rows_inlines and r_i < len(block.rows_inlines):
                        if c_i < len(block.rows_inlines[r_i]):
                            cell_inlines = block.rows_inlines[r_i][c_i]
                    cell_text = row[c_i] if c_i < len(row) else ""
                    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                    _clear_paragraph_runs(p)
                    if cell_inlines:
                        _apply_inlines(p, cell_inlines)
                    else:
                        p.add_run(cell_text)
                    if "TableText" in doc.styles:
                        p.style = doc.styles["TableText"]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _apply_three_line_table(table)
            continue

        if isinstance(block, CodeBlock):
            # 处理代码块
            p = doc.add_paragraph()
            if "CodeBlock" in doc.styles:
                p.style = doc.styles["CodeBlock"]
            elif "Body" in doc.styles:
                p.style = doc.styles["Body"]

            # 特殊处理 Mermaid 流程图
            if block.language and block.language.lower() == "mermaid":
                placeholder_run = p.add_run("[流程图占位：mermaid]")
                placeholder_run.italic = True
                # 添加代码内容作为参考
                p2 = doc.add_paragraph()
                if "CodeBlock" in doc.styles:
                    p2.style = doc.styles["CodeBlock"]
                code_run = p2.add_run(block.text or "")
                code_run.font.name = "Consolas"
            else:
                # 普通代码块
                code_run = p.add_run(block.text or "")
                code_run.font.name = "Consolas"
            continue

        if isinstance(block, FigureBlock):
            if spec.auto_number_figures_tables and block.caption and not re.match(r"^图\d+", block.caption.strip()):
                fig_counter += 1
                caption = f"图{fig_counter} {block.caption.strip()}"
            else:
                caption = block.caption
            if os.path.exists(block.path):
                doc.add_picture(block.path)
            else:
                p = doc.add_paragraph(f"[图片占位：{block.path}]")
                if "Body" in doc.styles:
                    p.style = doc.styles["Body"]
            if caption:
                pcap = doc.add_paragraph(caption)
                if "FigureCaption" in doc.styles:
                    pcap.style = doc.styles["FigureCaption"]
            continue

        if isinstance(block, PageBreakBlock):
            doc.add_page_break()
            continue

        if isinstance(block, SectionBreakBlock):
            doc.add_section(WD_SECTION.NEW_PAGE)
            continue

        if isinstance(block, BibliographyBlock):
            for it in block.items:
                p = doc.add_paragraph(it)
                if "Reference" in doc.styles:
                    p.style = doc.styles["Reference"]
            continue

    out = io.BytesIO()
    doc.save(out)
    data = out.getvalue()
    if need_page_numbering:
        _ensure_footer_page_numbers(doc, spec)
        out = io.BytesIO()
        doc.save(out)
        data = out.getvalue()
        data = _apply_page_numbering_ooxml(data, spec)
    return data


def _render_cover(doc: Document, ast: DocumentAST) -> None:
    if ast.meta.title_cn:
        p = doc.add_paragraph(ast.meta.title_cn)
        if "TitleCN" in doc.styles:
            p.style = doc.styles["TitleCN"]
    if ast.meta.title_en:
        p = doc.add_paragraph(ast.meta.title_en)
        if "TitleEN" in doc.styles:
            p.style = doc.styles["TitleEN"]

    meta_parts = []
    if ast.meta.major:
        meta_parts.append(f"专业：{ast.meta.major}")
    if ast.meta.author:
        meta_parts.append(f"学生：{ast.meta.author}")
    if ast.meta.tutor:
        meta_parts.append(f"指导教师：{ast.meta.tutor}")
    for line in meta_parts:
        p = doc.add_paragraph(line)
        if "MetaLine" in doc.styles:
            p.style = doc.styles["MetaLine"]


def _ensure_footer_page_numbers(doc: Document, spec: StyleSpec) -> None:
    pn = spec.page_numbering
    if not pn or not pn.enabled or not pn.show_in_footer:
        return

    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        for r in list(p.runs):
            try:
                p._p.remove(r._r)
            except Exception:
                pass
        if "PageNumber" in doc.styles:
            p.style = doc.styles["PageNumber"]
        p.alignment = _align_to_docx(pn.footer_alignment)
        run = p.add_run()
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), "PAGE")
        run._r.append(fld)


def _normalize_cn_keywords(txt: str) -> str:
    parts = [p for p in re.split(r"[，,;；\s]+", txt) if p]
    return "　".join(parts)


def _normalize_en_keywords(txt: str) -> str:
    parts = [p.strip() for p in re.split(r"[;；,，]+", txt) if p.strip()]
    return "; ".join(parts)


def _apply_three_line_table(table) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = tblPr.find(qn("w:tblBorders"))
    if tblBorders is None:
        tblBorders = OxmlElement("w:tblBorders")
        tblPr.append(tblBorders)

    def _border(tag: str, val: str, sz: int):
        el = tblBorders.find(qn(f"w:{tag}"))
        if el is None:
            el = OxmlElement(f"w:{tag}")
            tblBorders.append(el)
        el.set(qn("w:val"), val)
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")

    _border("top", "single", 12)
    _border("bottom", "single", 12)
    _border("insideH", "single", 6)
    _border("left", "nil", 0)
    _border("right", "nil", 0)
    _border("insideV", "nil", 0)
